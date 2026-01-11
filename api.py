from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import os
import tempfile
import docx
import time
import uuid

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.runnables import RunnableParallel, RunnablePassthrough
from langchain_core.messages import AIMessage, HumanMessage

app = FastAPI(title="RAG Chatbot API")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory storage for sessions
sessions = {}

class ChatMessage(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
    session_id: str
    question: str
    chat_history: List[ChatMessage] = []
    answer_type: str = "Short (2 Marks)"

class ChatResponse(BaseModel):
    answer: str
    session_id: str

class SessionResponse(BaseModel):
    session_id: str
    message: str

# ------------ EMBEDDINGS ------------
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# ------------ DOCX LOADER ------------
def load_docx(path):
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

# ------------ PROCESS DOCUMENTS ------------
def process_documents(files):
    docs = []

    for file in files:
        filename = file.filename
        ext = filename.split(".")[-1].lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            tmp.write(file.file.read())
            path = tmp.name

        try:
            if ext == "pdf":
                loader = PyPDFLoader(path)
                pdf_docs = loader.load()
                for p in pdf_docs:
                    if len(p.page_content.strip()) > 5:
                        docs.append(Document(page_content=p.page_content, metadata={"source": filename}))

            elif ext == "docx":
                text = load_docx(path)
                docs.append(Document(page_content=text, metadata={"source": filename}))

            elif ext == "txt":
                text = open(path, "r", encoding="utf-8", errors="ignore").read()
                docs.append(Document(page_content=text, metadata={"source": filename}))

        finally:
            os.remove(path)

    if not docs:
        return None

    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=150)
    chunks = splitter.split_documents(docs)

    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=f"rag_{int(time.time())}"
    )

    return vectordb

# ------------ BUILD RAG PIPELINE ------------
def build_rag(vectordb, answer_type="Short (2 Marks)"):
    retriever = vectordb.as_retriever(search_kwargs={"k": 4})

    # Customize prompt based on answer type
    answer_instructions = {
        "Short (2 Marks)": """
        Provide a brief, concise answer (2-3 sentences).
        Focus on the key point only.
        Suitable for 2-mark questions.
        """,
        "Medium (5 Marks)": """
        Provide a moderate length answer (1 paragraph, 5-7 sentences).
        Include main points with brief explanations.
        Suitable for 5-mark questions.
        """,
        "Detailed (10 Marks)": """
        Provide a comprehensive, detailed answer (2-3 paragraphs).
        Include definitions, explanations, examples, and important points.
        Use bullet points or numbered lists where appropriate.
        Suitable for 10-mark questions.
        """,
        "Viva/Interview": """
        Provide a SHORT, conversational answer (3-5 sentences maximum).
        Answer naturally as if speaking in an interview - be direct and to the point.
        Include ONE practical example or real-world application if relevant.
        Keep it brief and confident - viva answers should be spoken in 30-45 seconds.
        Do NOT write lengthy explanations.
        """
    }

    instruction = answer_instructions.get(answer_type, answer_instructions["Short (2 Marks)"])

    system_prompt = f"""
    You are a helpful study assistant.  
    ONLY use the provided context to answer.  
    If the answer is not found, say:
    "I don't know based on the provided material."

    ANSWER FORMAT:
    {instruction}

    Context:
    {{context}}
    """

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{question}")
    ])

    # Adjust temperature based on answer type
    temperature = 0.1 if answer_type == "Viva/Interview" else 0.2
    
    llm = ChatOllama(
        model="mistral", 
        temperature=temperature,
        base_url="http://localhost:11434"
    )

    def format_docs(docs):
        return "\n\n".join([d.page_content for d in docs])

    def create_chain_input(inputs):
        question = inputs if isinstance(inputs, str) else inputs.get("question", "")
        chat_history = inputs.get("chat_history", []) if isinstance(inputs, dict) else []
        
        # Retrieve and format context
        context_docs = retriever.invoke(question)
        context = format_docs(context_docs)
        
        return {
            "context": context,
            "question": question,
            "chat_history": chat_history
        }

    rag_chain = create_chain_input | prompt | llm

    return rag_chain

# ------------ API ENDPOINTS ------------

@app.get("/")
def read_root():
    return {"message": "RAG Chatbot API is running"}

@app.post("/upload", response_model=SessionResponse)
async def upload_files(files: List[UploadFile] = File(...)):
    """Upload and index documents"""
    try:
        vectordb = process_documents(files)
        if not vectordb:
            raise HTTPException(status_code=400, detail="No valid documents processed")
        
        session_id = str(uuid.uuid4())
        sessions[session_id] = {
            "vectordb": vectordb,
            "chat_history": []
        }
        
        return SessionResponse(
            session_id=session_id,
            message="Documents indexed successfully"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Chat with the RAG system"""
    if request.session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    try:
        session = sessions[request.session_id]
        vectordb = session["vectordb"]
        
        # Convert chat history to LangChain format
        chat_history = []
        for msg in request.chat_history:
            if msg.role == "user":
                chat_history.append(HumanMessage(content=msg.content))
            else:
                chat_history.append(AIMessage(content=msg.content))
        
        # Build and invoke RAG chain with answer type
        rag = build_rag(vectordb, request.answer_type)
        result = rag.invoke({
            "question": request.question,
            "chat_history": chat_history
        })
        
        # Extract content from result
        answer = result.content if hasattr(result, 'content') else str(result)
        
        # Update session history
        session["chat_history"].append({"role": "user", "content": request.question})
        session["chat_history"].append({"role": "assistant", "content": answer})
        
        return ChatResponse(
            answer=answer,
            session_id=request.session_id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/session/{session_id}")
async def delete_session(session_id: str):
    """Delete a session"""
    if session_id in sessions:
        del sessions[session_id]
        return {"message": "Session deleted"}
    raise HTTPException(status_code=404, detail="Session not found")

@app.get("/sessions")
async def list_sessions():
    """List all active sessions"""
    return {"sessions": list(sessions.keys())}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)